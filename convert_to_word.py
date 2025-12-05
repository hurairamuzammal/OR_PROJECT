from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

title = doc.add_heading('PP Chemicals - Operations Research Problem Solver', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Project Report', level=1)
doc.add_paragraph('')

doc.add_heading('Team Members', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'S.No'
hdr_cells[1].text = 'Roll Number'
hdr_cells[2].text = 'Name'
hdr_cells[3].text = 'Section'
for i in range(1, 5):
    row_cells = table.rows[i].cells
    row_cells[0].text = str(i)
    row_cells[1].text = '[Roll Number]'
    row_cells[2].text = '[Name]'
    row_cells[3].text = '[Section]'

doc.add_page_break()

doc.add_heading('1. Problem Statement', level=1)
doc.add_paragraph('PP Chemicals is a chemical manufacturing company that makes products for roads and buildings in Pakistan. The company faces three main problems in its daily operations:')

doc.add_heading('1.1 Production Planning Problem', level=2)
doc.add_paragraph('The company makes 10 different chemical products like Bitumen Emulsion, Modified Bitumen, Concrete Plasticizer, and others. Each product needs different amounts of raw materials, machine time, labor hours, and energy. The company wants to know how much of each product to make so that they can get the most profit while using the resources they have.')

doc.add_heading('1.2 Worker Assignment Problem', level=2)
doc.add_paragraph('The company has 10 workers and 10 different tasks like Mixing, Heating, Testing, Packing, and Quality Control. Each worker has different skills and works at different speeds on different tasks. The company wants to assign each worker to one task in a way that the total work done is the best possible.')

doc.add_heading('1.3 Transportation Problem', level=2)
doc.add_paragraph('The company has 10 plants in different cities of Pakistan like Karachi, Lahore, Islamabad, and others. They need to send their products to 10 construction sites like M-2 Motorway, Lahore Metro, and Gwadar Port. Each route has a different cost to ship products. The company wants to find the cheapest way to send products from all plants to all sites while meeting the supply and demand needs.')

doc.add_heading('2. Objectives', level=1)
objectives = [
    'Build a desktop application that can solve these three types of problems',
    'Help managers make better decisions by showing them the best solutions',
    'Provide sensitivity analysis to see how changes in data affect the solutions',
    'Create easy-to-use screens where users can enter their data and see results',
    'Support large problems with 10 or more variables and constraints'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Number')

doc.add_page_break()

doc.add_heading('3. Formulation of the Problems', level=1)

doc.add_heading('3.1 Linear Programming Problem (Simplex Method)', level=2)

doc.add_heading('Decision Variables:', level=3)
doc.add_paragraph('Let x₁, x₂, x₃, ..., x₁₀ be the amount (in tons) of each product to make:')
variables = ['x₁ = Bitumen Emulsion', 'x₂ = Modified Bitumen', 'x₃ = Concrete Plasticizer', 
             'x₄ = Curing Compound', 'x₅ = Waterproofing Compound', 'x₆ = Road Marking Paint',
             'x₇ = Anti-Strip Agent', 'x₈ = Concrete Hardener', 'x₉ = Epoxy Coating', 
             'x₁₀ = Polymer Modified Bitumen']
for var in variables:
    doc.add_paragraph(var, style='List Bullet')

doc.add_heading('Objective Function:', level=3)
doc.add_paragraph('Maximize Z = 5000x₁ + 7500x₂ + 4000x₃ + 3500x₄ + 6000x₅ + 4500x₆ + 8000x₇ + 3000x₈ + 9000x₉ + 8500x₁₀')
doc.add_paragraph('(where the numbers are profit in Rs. per ton of each product)')

doc.add_heading('Constraints:', level=3)
doc.add_paragraph('The problem has 10 constraints for different resources:')
constraints = [
    'Raw Material A: 2x₁ + 3x₂ + 1x₃ + 2x₄ + 1x₅ + 2x₆ + 3x₇ + 1x₈ + 2x₉ + 3x₁₀ ≤ 5000 kg',
    'Raw Material B: 1x₁ + 2x₂ + 3x₃ + 1x₄ + 2x₅ + 1x₆ + 2x₇ + 3x₈ + 1x₉ + 2x₁₀ ≤ 4000 kg',
    'Production Line 1: 3x₁ + 2x₂ + 4x₃ + 1x₄ + 2x₅ + 3x₆ + 1x₇ + 2x₈ + 4x₉ + 2x₁₀ ≤ 480 hours',
    'Production Line 2: 1x₁ + 2x₂ + 1x₃ + 3x₄ + 4x₅ + 2x₆ + 1x₇ + 2x₈ + 1x₉ + 3x₁₀ ≤ 400 hours',
    'Storage Capacity: 1x₁ + 1x₂ + 1x₃ + 1x₄ + 1x₅ + 1x₆ + 1x₇ + 1x₈ + 1x₉ + 1x₁₀ ≤ 1000 tons',
    'Labor Hours: 4x₁ + 5x₂ + 3x₃ + 4x₄ + 5x₅ + 3x₆ + 4x₇ + 5x₈ + 3x₉ + 4x₁₀ ≤ 2000 man-hours',
    'Quality Control: 0.5x₁ + 0.5x₂ + 0.5x₃ + 0.5x₄ + 0.5x₅ + 0.5x₆ + 0.5x₇ + 0.5x₈ + 0.5x₉ + 0.5x₁₀ ≤ 300 tests',
    'Environmental Limit: 0.1x₁ + 0.2x₂ + 0.15x₃ + 0.1x₄ + 0.2x₅ + 0.15x₆ + 0.1x₇ + 0.2x₈ + 0.15x₉ + 0.1x₁₀ ≤ 100 units',
    'Energy: 10x₁ + 15x₂ + 8x₃ + 12x₄ + 10x₅ + 8x₆ + 15x₇ + 10x₈ + 12x₉ + 14x₁₀ ≤ 10000 kWh',
    'Packaging Material: 1x₁ + 2x₂ + 1x₃ + 1x₄ + 2x₅ + 1x₆ + 2x₇ + 1x₈ + 1x₉ + 2x₁₀ ≤ 2500 units'
]
for i, con in enumerate(constraints, 1):
    doc.add_paragraph(f'{i}. {con}')

doc.add_heading('Non-negativity:', level=3)
doc.add_paragraph('x₁, x₂, x₃, x₄, x₅, x₆, x₇, x₈, x₉, x₁₀ ≥ 0')

doc.add_heading('How Data was Generated:', level=3)
doc.add_paragraph('The profit values are based on market prices for chemical products in Pakistan. The resource needs are based on typical production needs in chemical factories. The available amounts of resources are based on a medium-sized factory capacity.')

doc.add_heading('Solution Method:', level=3)
doc.add_paragraph('We use the Simplex Method with the HiGHS solver from scipy library. This method finds the best solution step by step until no more improvement is possible.')

doc.add_page_break()

doc.add_heading('3.2 Assignment Problem (Hungarian Algorithm)', level=2)

doc.add_heading('Problem Setup:', level=3)
doc.add_paragraph('We have a 10×10 matrix where rows are workers and columns are tasks. Each cell shows how good a worker is at that task (efficiency score from 0 to 100).')

doc.add_heading('Workers:', level=3)
workers = ['Ali Khan', 'Bilal Ahmed', 'Chaudhry Imran', 'Danish Malik', 'Ejaz Shah',
           'Farhan Raza', 'Ghulam Abbas', 'Hassan Javed', 'Irfan Siddiqui', 'Junaid Tariq']
for i, w in enumerate(workers, 1):
    doc.add_paragraph(f'{i}. {w}')

doc.add_heading('Tasks:', level=3)
tasks = ['Mixing', 'Heating', 'Testing', 'Packing', 'Loading',
         'Quality Control', 'Maintenance', 'Documentation', 'Safety Check', 'Dispatch']
for i, t in enumerate(tasks, 1):
    doc.add_paragraph(f'{i}. {t}')

doc.add_heading('Efficiency Matrix:', level=3)
efficiency_data = [
    ['', 'Mix', 'Heat', 'Test', 'Pack', 'Load', 'QC', 'Maint', 'Doc', 'Safety', 'Disp'],
    ['Ali', '85', '70', '65', '80', '75', '90', '60', '50', '70', '65'],
    ['Bilal', '75', '85', '70', '65', '80', '75', '70', '60', '65', '70'],
    ['Chaudhry', '80', '75', '90', '70', '65', '80', '75', '55', '80', '60'],
    ['Danish', '65', '80', '75', '90', '70', '65', '80', '70', '75', '80'],
    ['Ejaz', '70', '65', '80', '75', '90', '70', '65', '75', '60', '85'],
    ['Farhan', '90', '75', '70', '65', '80', '85', '70', '80', '65', '70'],
    ['Ghulam', '60', '90', '65', '80', '75', '70', '95', '65', '80', '75'],
    ['Hassan', '55', '65', '80', '70', '85', '75', '70', '90', '70', '80'],
    ['Irfan', '75', '70', '85', '75', '70', '80', '75', '70', '95', '65'],
    ['Junaid', '70', '75', '60', '85', '80', '70', '80', '75', '70', '90']
]
eff_table = doc.add_table(rows=11, cols=11)
eff_table.style = 'Table Grid'
for i, row_data in enumerate(efficiency_data):
    for j, cell_data in enumerate(row_data):
        eff_table.rows[i].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_heading('Objective:', level=3)
doc.add_paragraph('Maximize Total Efficiency = Sum of efficiency scores of all assignments')

doc.add_heading('Constraint:', level=3)
doc.add_paragraph('Each worker must be assigned to exactly one task, and each task must have exactly one worker.')

doc.add_heading('Solution Method:', level=3)
doc.add_paragraph('We use the Hungarian Algorithm from scipy library. This algorithm works in O(n³) time and always finds the best assignment.')

doc.add_page_break()

doc.add_heading('3.3 Transportation Problem (VAM + MODI Method)', level=2)

doc.add_heading('Problem Setup:', level=3)
doc.add_paragraph('We have 10 source plants and 10 destination sites. Each plant has a supply amount, each site has a demand amount, and each route has a shipping cost.')

doc.add_heading('Sources (Plants) with Supply (tons/month):', level=3)
sources = [('Karachi', 500), ('Lahore', 400), ('Islamabad', 350), ('Faisalabad', 450), ('Rawalpindi', 380),
           ('Multan', 420), ('Peshawar', 300), ('Quetta', 360), ('Sialkot', 410), ('Gujranwala', 330)]
for i, (name, supply) in enumerate(sources, 1):
    doc.add_paragraph(f'{i}. {name} - {supply}')

doc.add_heading('Destinations (Sites) with Demand (tons/month):', level=3)
dests = [('M-2 Motorway', 200), ('Lahore Metro', 180), ('Attock Bridge', 300), ('Karachi Flyover', 250), ('GT Road', 350),
         ('Lowari Tunnel', 280), ('Islamabad Airport', 320), ('Gwadar Port', 400), ('Peshawar Railway', 290), ('Multan Stadium', 330)]
for i, (name, demand) in enumerate(dests, 1):
    doc.add_paragraph(f'{i}. {name} - {demand}')

doc.add_heading('Cost Matrix (Rs. per ton):', level=3)
cost_data = [
    ['', 'M-2', 'Metro', 'Attock', 'Karachi', 'GT', 'Lowari', 'Airport', 'Gwadar', 'Pesh', 'Multan'],
    ['Karachi', '45', '72', '35', '58', '62', '48', '55', '80', '42', '65'],
    ['Lahore', '38', '65', '42', '52', '58', '45', '50', '75', '38', '60'],
    ['Islamabad', '55', '48', '58', '42', '45', '52', '48', '62', '55', '45'],
    ['Faisalabad', '62', '55', '48', '38', '42', '55', '52', '58', '48', '42'],
    ['Rawalpindi', '70', '58', '52', '45', '38', '48', '45', '52', '55', '48'],
    ['Multan', '58', '52', '55', '48', '45', '35', '42', '48', '52', '55'],
    ['Peshawar', '85', '78', '72', '65', '58', '52', '45', '38', '65', '58'],
    ['Quetta', '78', '72', '65', '58', '52', '48', '42', '45', '58', '55'],
    ['Sialkot', '72', '68', '62', '55', '48', '52', '48', '42', '52', '48'],
    ['Gujranwala', '95', '88', '82', '75', '68', '62', '55', '48', '72', '65']
]
cost_table = doc.add_table(rows=11, cols=11)
cost_table.style = 'Table Grid'
for i, row_data in enumerate(cost_data):
    for j, cell_data in enumerate(row_data):
        cost_table.rows[i].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_heading('Objective:', level=3)
doc.add_paragraph('Minimize Total Transportation Cost = Sum of (quantity shipped × cost per ton) for all routes')

doc.add_heading('Constraints:', level=3)
doc.add_paragraph('1. Total shipped from each source = Supply at that source')
doc.add_paragraph('2. Total received at each destination = Demand at that destination')

doc.add_heading('Solution Method:', level=3)
doc.add_paragraph('1. Initial Solution: We use Vogel\'s Approximation Method (VAM) to find a starting solution. VAM looks at penalty costs and picks routes with the biggest penalty difference first.')
doc.add_paragraph('2. Optimization: We use MODI (Modified Distribution) Method to improve the initial solution. MODI calculates opportunity costs and keeps improving until no better solution exists.')

doc.add_page_break()

doc.add_heading('4. Results', level=1)

doc.add_heading('4.1 Linear Programming Results', level=2)

doc.add_heading('Dashboard Screen:', level=3)
doc.add_paragraph('The main screen shows three cards for the three problem types. Users can click on any card to open that solver.')

doc.add_heading('Simplex Input Screen:', level=3)
doc.add_paragraph('Users can enter:')
inputs = ['Number of variables (products) and constraints (resources)', 
          'Objective function coefficients (profit per product)',
          'Constraint matrix (resources needed per product)',
          'Right-hand side values (available resources)',
          'Choose maximize or minimize']
for inp in inputs:
    doc.add_paragraph(inp, style='List Bullet')

doc.add_heading('Solution Display:', level=3)
doc.add_paragraph('After clicking "Solve Problem", the system shows:')
outputs = ['Status: Optimal solution found',
           'Optimal Profit: The maximum profit value',
           'Production Plan: How much of each product to make',
           'Iterations: How many steps the solver took']
for out in outputs:
    doc.add_paragraph(out, style='List Bullet')

doc.add_heading('Sensitivity Analysis:', level=3)
doc.add_paragraph('The system shows:')
sens = ['Shadow Prices: How much profit increases if we get one more unit of each resource',
        'Reduced Costs: How much a product\'s profit must increase before we should make it',
        'Slack Values: How much of each resource is left unused',
        'Allowable Ranges: How much parameters can change without changing the basic solution']
for s in sens:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Sensitivity Interpretation:', level=3)
doc.add_paragraph('If a shadow price is Rs. 1500 for Raw Material A, it means adding 1 more kg of this material will increase profit by Rs. 1500')
doc.add_paragraph('If a slack value is 0, that resource is fully used (binding constraint)')
doc.add_paragraph('If allowable increase is infinity, we can add any amount without changing the solution structure')

doc.add_heading('4.2 Assignment Problem Results', level=2)

doc.add_heading('Sample Output:', level=3)
assignments = [
    'Ali Khan → Quality Control (Efficiency: 90)',
    'Bilal Ahmed → Heating (Efficiency: 85)',
    'Chaudhry Imran → Testing (Efficiency: 90)',
    'Danish Malik → Packing (Efficiency: 90)',
    'Ejaz Shah → Loading (Efficiency: 90)',
    'Farhan Raza → Mixing (Efficiency: 90)',
    'Ghulam Abbas → Maintenance (Efficiency: 95)',
    'Hassan Javed → Documentation (Efficiency: 90)',
    'Irfan Siddiqui → Safety Check (Efficiency: 95)',
    'Junaid Tariq → Dispatch (Efficiency: 90)'
]
doc.add_paragraph('Optimal Assignments:')
for a in assignments:
    doc.add_paragraph(a, style='List Bullet')
doc.add_paragraph('Total Efficiency Score: 905', style='Intense Quote')

doc.add_heading('4.3 Transportation Problem Results', level=2)
doc.add_paragraph('The system shows:')
trans_results = ['Total Transportation Cost: Minimum cost to ship all products',
                 'Initial Method Used: Which method was used for starting solution',
                 'MODI Iterations: How many improvement steps were needed',
                 'Route Details: From which plant to which site, how much, and at what cost']
for tr in trans_results:
    doc.add_paragraph(tr, style='List Bullet')

# Section 4.4 - Screenshots
doc.add_page_break()
doc.add_heading('4.4 Application Screenshots', level=2)

import os
screenshots_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'screenshots')

# 4.4.1 Linear Programming - Problem Input
doc.add_heading('4.4.1 Linear Programming - Problem Input', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'lp_input.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: lp_input.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
lp_input_points = [
    'Problem Inputs panel for Linear Programming',
    'Users can set number of variables and constraints',
    'Objective function with profit/cost for each product',
    'Products: Bitumen Emulsion, Modified Bitumen, Concrete Primer, etc.',
    'Option to Maximize Profit or Minimize Cost'
]
for point in lp_input_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.2 Linear Programming - Solution Output
doc.add_heading('4.4.2 Linear Programming - Solution Output', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'lp_solution.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: lp_solution.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
lp_solution_points = [
    'Results & Analysis panel with solved solution',
    'Optimal Value: Rs. 3,226,666.67',
    'Number of iterations: 6',
    'Decision variables showing how much to produce',
    'Each product\'s recommended quantity'
]
for point in lp_solution_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.3 Linear Programming - Fullscreen Results
doc.add_heading('4.4.3 Linear Programming - Fullscreen Results', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'results_fullscreen.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: results_fullscreen.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
lp_fullscreen_points = [
    'Fullscreen view of complete results',
    'Optimal Production Plan:',
    '  - Anti-Strip Agent: 373.33 units',
    '  - Epoxy Coating: 26.67 units',
    'Sensitivity Analysis with shadow prices:',
    '  - Raw Material A: Rs. 0.00',
    '  - Production Line 1: Rs. 333.33',
    '  - Production Line 2: Rs. 7,666.67'
]
for point in lp_fullscreen_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.4 Assignment Problem - Cost/Efficiency Matrix
doc.add_heading('4.4.4 Assignment Problem - Cost/Efficiency Matrix', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'assignment_matrix.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: assignment_matrix.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
assign_matrix_points = [
    '10x10 Cost/Efficiency Matrix in fullscreen',
    'Rows = Workers (Ali Khan, Bilal Ahmed, etc.)',
    'Columns = Tasks (Mixing, Heating, Testing, etc.)',
    'Each cell = efficiency score (0 to 100)',
    'Higher score means worker is better at that task'
]
for point in assign_matrix_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.5 Assignment Problem - Optimal Assignments
doc.add_heading('4.4.5 Assignment Problem - Optimal Assignments', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'assignment_results.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: assignment_results.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
assign_results_points = [
    'Optimal Total Efficiency: 905.00',
    'Each worker assigned to their best task:',
    '  - Ali Khan → QC (90)',
    '  - Bilal Ahmed → Heating (85)',
    '  - Chaudhry Imran → Testing (90)',
    '  - Ghulam Abbas → Maintenance (95)',
    '  - Irfan Siddiqui → Safety (95)'
]
for point in assign_results_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.6 Assignment Problem - Full Interface
doc.add_heading('4.4.6 Assignment Problem - Full Interface', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'assignment_interface.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: assignment_interface.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
assign_interface_points = [
    'Left panel: Problem Inputs with Cost/Efficiency Matrix',
    'Right panel: Results showing optimal assignments',
    'Buttons: Find Optimal Assignment, Clear All, Random Data',
    'Status: Optimal Assignment Found'
]
for point in assign_interface_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.7 Transportation Problem - Cost Matrix
doc.add_heading('4.4.7 Transportation Problem - Cost Matrix', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'transportation_matrix.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: transportation_matrix.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
trans_matrix_points = [
    '10x10 Transportation Cost Matrix',
    'Sources: Karachi, Lahore, Islamabad, Faisalabad, etc.',
    'Destinations: M-2 Motorway, Lahore Metro, GT Road, etc.',
    'Total Supply: 3,900 units',
    'Total Demand: 2,900 units',
    'Problem is unbalanced (supply > demand)'
]
for point in trans_matrix_points:
    doc.add_paragraph(point, style='List Bullet')

# 4.4.8 Transportation Problem - Optimal Routes
doc.add_heading('4.4.8 Transportation Problem - Optimal Routes', level=3)
try:
    doc.add_picture(os.path.join(screenshots_dir, 'transportation_results.png'), width=Inches(5.5))
except:
    doc.add_paragraph('[Image: transportation_results.png]')
doc.add_paragraph('What this shows:', style='Intense Quote')
trans_routes_points = [
    'Total Transportation Cost: Rs. 115,480.00',
    'Optimal Shipping Routes list',
    'Shows source plant and destination site',
    'Cost per unit and quantity shipped',
    'Solution uses VAM + MODI method'
]
for point in trans_routes_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_page_break()

doc.add_heading('5. Codes', level=1)

doc.add_heading('5.1 Simplex Solver Code (simplex.py)', level=2)
simplex_code = '''import numpy as np
from scipy.optimize import linprog
from dataclasses import dataclass
from typing import Optional, List, Tuple, Dict, Any

@dataclass
class SensitivityAnalysis:
    shadow_prices: np.ndarray
    reduced_costs: np.ndarray
    slack_values: np.ndarray
    constraint_rhs_ranges: List[Dict]
    objective_coeff_ranges: List[Dict]

@dataclass
class SimplexResult:
    success: bool
    message: str
    optimal_value: float
    solution: np.ndarray
    sensitivity: Optional[SensitivityAnalysis]
    iterations: int
    status: int

class SimplexSolver:
    def __init__(self, c, A_ub=None, b_ub=None, maximize=True):
        self.c = np.array(c, dtype=float)
        self.A_ub = np.array(A_ub, dtype=float) if A_ub is not None else None
        self.b_ub = np.array(b_ub, dtype=float) if b_ub is not None else None
        self.maximize = maximize

    def solve(self):
        c_solve = -self.c if self.maximize else self.c
        result = linprog(c_solve, A_ub=self.A_ub, b_ub=self.b_ub, method='highs')
        if result.success:
            optimal_value = -result.fun if self.maximize else result.fun
            return SimplexResult(True, "Optimal", optimal_value, result.x, None, result.nit, 0)
        return SimplexResult(False, result.message, 0.0, None, None, 0, -1)'''
code_para = doc.add_paragraph()
code_run = code_para.add_run(simplex_code)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(8)

doc.add_heading('5.2 Assignment Solver Code (assignment.py)', level=2)
assign_code = '''import numpy as np
from scipy.optimize import linear_sum_assignment
from dataclasses import dataclass
from typing import List, Tuple

@dataclass
class AssignmentResult:
    success: bool
    total_cost: float
    assignments: List[Tuple[int, int]]
    assignment_matrix: np.ndarray

class AssignmentSolver:
    def __init__(self, cost_matrix, maximize=False):
        self.cost_matrix = np.array(cost_matrix, dtype=float)
        self.maximize = maximize

    def solve(self):
        if self.maximize:
            solve_matrix = -self.cost_matrix
        else:
            solve_matrix = self.cost_matrix
        row_ind, col_ind = linear_sum_assignment(solve_matrix)
        assignments = list(zip(row_ind, col_ind))
        total = sum(self.cost_matrix[r, c] for r, c in assignments)
        matrix = np.zeros_like(self.cost_matrix, dtype=int)
        for r, c in assignments:
            matrix[r, c] = 1
        return AssignmentResult(True, total, assignments, matrix)'''
code_para2 = doc.add_paragraph()
code_run2 = code_para2.add_run(assign_code)
code_run2.font.name = 'Consolas'
code_run2.font.size = Pt(8)

doc.add_heading('5.3 Transportation Solver Code (transportation.py)', level=2)
trans_code = '''import numpy as np
from dataclasses import dataclass
from typing import List, Dict
from enum import Enum

class InitialMethod(Enum):
    NORTH_WEST_CORNER = "north_west"
    LEAST_COST = "least_cost"
    VOGEL = "vam"

@dataclass
class TransportationResult:
    success: bool
    total_cost: float
    allocation_matrix: np.ndarray
    route_details: List[Dict]

class TransportationSolver:
    def __init__(self, supply, demand, cost_matrix):
        self.supply = np.array(supply, dtype=float)
        self.demand = np.array(demand, dtype=float)
        self.cost_matrix = np.array(cost_matrix, dtype=float)
        self.m = len(supply)
        self.n = len(demand)
        self._balance_problem()

    def _balance_problem(self):
        total_supply = np.sum(self.supply)
        total_demand = np.sum(self.demand)
        if total_supply > total_demand:
            self.demand = np.append(self.demand, total_supply - total_demand)
            self.cost_matrix = np.hstack([self.cost_matrix, np.zeros((self.m, 1))])
            self.n += 1

    def solve(self, method=InitialMethod.VOGEL):
        allocation = self._vogel_approximation_method()
        total_cost = np.sum(allocation * self.cost_matrix)
        return TransportationResult(True, total_cost, allocation, [])'''
code_para3 = doc.add_paragraph()
code_run3 = code_para3.add_run(trans_code)
code_run3.font.name = 'Consolas'
code_run3.font.size = Pt(8)

doc.add_heading('5.4 Main Application (main.py)', level=2)
main_code = '''import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from ui.app import App

def main():
    print("PP CHEMICALS - OR Problem Solver")
    print("Starting application...")
    try:
        app = App()
        app.mainloop()
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()'''
code_para4 = doc.add_paragraph()
code_run4 = code_para4.add_run(main_code)
code_run4.font.name = 'Consolas'
code_run4.font.size = Pt(8)

doc.add_page_break()

doc.add_heading('6. Conclusion', level=1)
conclusion = '''In this project, we built a desktop application to solve three important Operations Research problems for PP Chemicals company. The application uses well-known algorithms to find the best solutions:

1. Linear Programming: The Simplex Method helps the company find the best production plan to maximize profit. With 10 products and 10 resource types, the solver finds how much of each product to make. The sensitivity analysis tells managers which resources are most valuable and how much changes in data affect the solution.

2. Assignment Problem: The Hungarian Algorithm finds the best way to assign 10 workers to 10 tasks. By maximizing the total efficiency score (905 in our sample), the company can make sure each worker does the task they are best at.

3. Transportation Problem: Using VAM for the initial solution and MODI for optimization, we find the cheapest way to ship products from 10 plants to 10 construction sites. The solution shows exactly how many tons to ship on each route to minimize the total shipping cost.

The application has an easy-to-use interface built with CustomTkinter. Users can enter their data, load sample problems, and see results with clear displays. The code is organized into separate files for algorithms, user interface, and settings, making it easy to understand and change.

This tool can help PP Chemicals managers make better decisions about production, worker assignment, and shipping. The sensitivity analysis features are especially useful because they show how the solutions change when the data changes, helping managers plan for different situations.'''
doc.add_paragraph(conclusion)

doc.add_paragraph('')
end_para = doc.add_paragraph('*********************END********************')
end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save('PROJECT_DOCUMENTATION.docx')
print("Word document created: PROJECT_DOCUMENTATION.docx")
